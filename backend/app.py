import os
import requests
import bcrypt
import jwt
from pymongo import MongoClient
from datetime import datetime, timedelta
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from dotenv import load_dotenv
from functools import wraps
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import json

# Heavy ML libraries will be lazy-imported to allow fast app startup
# Globals to hold loaded models/pipelines
model = None
summarizer = None
nlp = None

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv("SECRET_KEY", "your-secret-key-change-in-production")
CORS(app)

# MongoDB connection
MONGO_URL = os.getenv("MONGO_URL", "mongodb://localhost:27017")
client = MongoClient(MONGO_URL)
db = client["meeting_minutes"]
users_collection = db["users"]
transcriptions_collection = db["transcriptions"]
logs_collection = db["logs"]
analytics_collection = db["analytics"]

# Check MongoDB availability; if not available, fall back to in-memory stores for development
USE_MONGO = True
try:
    # quick ping
    client.admin.command('ping')
except Exception:
    print("Warning: Could not connect to MongoDB at %s; falling back to in-memory stores for development." % MONGO_URL)
    USE_MONGO = False
    # Simple in-memory fallbacks (not persistent)
    _in_memory_id = 1
    fallback_users = {}
    fallback_transcriptions = {}

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Note: heavy ML models (Whisper, Transformers, spaCy) are loaded on-demand

# ===========================
# LOGGING UTILITY
# ===========================

def log_action(action, user_id, details=None):
    """Log user actions to MongoDB"""
    log_entry = {
        "action": action,
        "user_id": str(user_id) if user_id else None,
        "timestamp": datetime.utcnow(),
        "details": details or {}
    }
    logs_collection.insert_one(log_entry)

def track_metric(metric_type, value, user_id=None):
    """Track metrics for analytics"""
    metric = {
        "type": metric_type,
        "value": value,
        "user_id": str(user_id) if user_id else None,
        "timestamp": datetime.utcnow()
    }
    analytics_collection.insert_one(metric)

# ===========================
# AUTHENTICATION MIDDLEWARE
# ===========================

def token_required(f):
    """Decorator to protect routes that require authentication"""
    @wraps(f)
    def decorated(*args, **kwargs):
        token = None
        
        # Check for token in headers
        if 'Authorization' in request.headers:
            auth_header = request.headers['Authorization']
            try:
                token = auth_header.split(" ")[1]
            except IndexError:
                return jsonify({"error": "Invalid token format"}), 401
        
        if not token:
            return jsonify({"error": "Token is missing"}), 401
        
        try:
            data = jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
            current_user_id = data['user_id']
        except jwt.ExpiredSignatureError:
            return jsonify({"error": "Token has expired"}), 401
        except jwt.InvalidTokenError:
            return jsonify({"error": "Invalid token"}), 401
        
        return f(current_user_id, *args, **kwargs)
    
    return decorated

def admin_required(f):
    """Decorator to check if user is admin"""
    @wraps(f)
    def decorated(*args, **kwargs):
        token = None
        
        if 'Authorization' in request.headers:
            auth_header = request.headers['Authorization']
            try:
                token = auth_header.split(" ")[1]
            except IndexError:
                return jsonify({"error": "Invalid token format"}), 401
        
        if not token:
            return jsonify({"error": "Token is missing"}), 401
        
        try:
            data = jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
            current_user_id = data['user_id']
            
            # Check if user is admin
            user = None
            if USE_MONGO:
                try:
                    from bson import ObjectId
                    user = users_collection.find_one({"_id": ObjectId(current_user_id)})
                except Exception:
                    user = None

            if not user:
                user = fallback_users.get(str(current_user_id))

            if not user or user.get("role") != "admin":
                return jsonify({"error": "Admin access required"}), 403
                
        except jwt.ExpiredSignatureError:
            return jsonify({"error": "Token has expired"}), 401
        except jwt.InvalidTokenError:
            return jsonify({"error": "Invalid token"}), 401
        
        return f(current_user_id, *args, **kwargs)
    
    return decorated

# ===========================
# AUDIO PROCESSING UTILITIES
# ===========================

def denoise_audio(filepath):
    """Remove noise from audio file"""
    try:
        # Lazy-import heavy audio libs
        import librosa
        import noisereduce as nr
        import soundfile as sf

        # Load audio
        y, sr = librosa.load(filepath, sr=None)

        # Denoise
        y_denoised = nr.reduce_noise(y=y, sr=sr)

        # Save denoised audio
        denoised_path = filepath.replace(".wav", "_denoised.wav").replace(".mp3", "_denoised.wav")
        sf.write(denoised_path, y_denoised, sr)

        return denoised_path
    except Exception as e:
        print(f"Denoising failed or audio libs not available: {e}")
        return filepath

def extract_segments(result):
    """Extract segments with timestamps from Whisper result"""
    segments = []
    if "segments" in result:
        for segment in result["segments"]:
            segments.append({
                "start": segment.get("start"),
                "end": segment.get("end"),
                "text": segment.get("text")
            })
    return segments

# ===========================
# KEY ITEMS EXTRACTION (AD-6)
# ===========================
try:
    import spacy
    try:
        nlp = spacy.load("en_core_web_sm")
    except Exception:
        nlp = None
        print("spaCy model en_core_web_sm not loaded; run 'python -m spacy download en_core_web_sm' to enable key item extraction")
except Exception:
    nlp = None
    print("spaCy not installed; key item extraction disabled")


def extract_key_items_from_text(text):
    """Simple rule-based extraction of action items and decisions using spaCy when available."""
    items = []
    global nlp
    # Try to initialize spaCy on-demand
    if not nlp:
        try:
            import spacy
            try:
                nlp = spacy.load("en_core_web_sm")
            except Exception:
                nlp = None
                print("spaCy model en_core_web_sm not loaded; run 'python -m spacy download en_core_web_sm' to enable key item extraction")
        except Exception:
            nlp = None
            print("spaCy not installed; key item extraction disabled")

    if not nlp:
        return items

    doc = nlp(text)
    for sent in doc.sents:
        s_text = sent.text.strip()
        lowered = s_text.lower()
        # heuristics for action items / decisions
        if any(keyword in lowered for keyword in ["action:", "action item", "todo", "to do", "will", "should", "agree", "decide", "decision"]):
            # try to extract assignee and task
            assignee = None
            for ent in sent.ents:
                if ent.label_ in ("PERSON", "ORG"):
                    assignee = ent.text
                    break
            items.append({"text": s_text, "assignee": assignee, "status": "open"})
        else:
            # pattern: "<Person> will <verb> ..."
            for token in sent:
                if token.lemma_ == "will":
                    # find nearest person entity
                    assignee = None
                    for ent in sent.ents:
                        if ent.label_ == "PERSON":
                            assignee = ent.text
                            break
                    items.append({"text": s_text, "assignee": assignee, "status": "open"})
                    break

    return items



# ===========================
# AUTH ENDPOINTS
# ===========================

@app.route("/register", methods=["POST"])
def register():
    """Register a new user"""
    data = request.get_json()
    
    # Validation
    if not data or not data.get("email") or not data.get("password") or not data.get("name"):
        return jsonify({"error": "Email, name, and password are required"}), 400
    
    email = data.get("email").lower()
    password = data.get("password")
    name = data.get("name")
    
    # Check if user already exists
    if USE_MONGO:
        try:
            if users_collection.find_one({"email": email}):
                return jsonify({"error": "Email already registered"}), 409
        except Exception:
            # fall back to in-memory on DB error
            pass
    else:
        if any(u.get("email") == email for u in fallback_users.values()):
            return jsonify({"error": "Email already registered"}), 409
    
    # Validate password length
    if len(password) < 6:
        return jsonify({"error": "Password must be at least 6 characters"}), 400
    
    # Hash password
    hashed_password = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())
    
    # Create user
    user = {
        "name": name,
        "email": email,
        "password": hashed_password,
        "created_at": datetime.utcnow(),
        "updated_at": datetime.utcnow()
    }
    
    if USE_MONGO:
        try:
            result = users_collection.insert_one(user)
            user_id = str(result.inserted_id)
        except Exception:
            # fallback to in-memory
            im = globals().get("_in_memory_id", 1)
            user_id = str(im)
            fallback_users[user_id] = user
            globals()["_in_memory_id"] = im + 1
    else:
        im = globals().get("_in_memory_id", 1)
        user_id = str(im)
        fallback_users[user_id] = user
        globals()["_in_memory_id"] = im + 1

    return jsonify({
        "message": "User registered successfully",
        "user_id": str(user_id)
    }), 201


@app.route("/login", methods=["POST"])
def login():
    """Login user and return JWT token"""
    data = request.get_json()
    
    if not data or not data.get("email") or not data.get("password"):
        return jsonify({"error": "Email and password are required"}), 400
    
    email = data.get("email").lower()
    password = data.get("password")
    
    # Find user
    user = None
    if USE_MONGO:
        try:
            user = users_collection.find_one({"email": email})
        except Exception:
            user = None

    if not user:
        # check fallback store
        for uid, u in (fallback_users.items() if not USE_MONGO else fallback_users.items()):
            if u.get("email") == email:
                user = u
                user['_id'] = uid
                break

    if not user:
        return jsonify({"error": "Invalid email or password"}), 401

    # Check password
    stored_pw = user.get('password')
    if isinstance(stored_pw, str):
        # if somehow stored as str, convert
        stored_pw = stored_pw.encode('utf-8')

    if not bcrypt.checkpw(password.encode('utf-8'), stored_pw):
        return jsonify({"error": "Invalid email or password"}), 401
    
    # Generate JWT token
    token = jwt.encode({
        'user_id': str(user['_id']),
        'email': user['email'],
        'exp': datetime.utcnow() + timedelta(hours=24)
    }, app.config['SECRET_KEY'], algorithm='HS256')
    
    return jsonify({
        "message": "Login successful",
        "token": token,
        "user": {
            "id": str(user['_id']),
            "name": user['name'],
            "email": user['email']
        }
    }), 200


@app.route("/verify-token", methods=["POST"])
def verify_token():
    """Verify if a token is valid"""
    data = request.get_json()
    token = data.get("token")
    
    if not token:
        return jsonify({"error": "Token is required"}), 400
    
    try:
        data = jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
        user_id = data['user_id']

        # Fetch user record to include role and name (DB or fallback)
        user = None
        if USE_MONGO:
            try:
                from bson import ObjectId
                user = users_collection.find_one({"_id": ObjectId(user_id)})
            except Exception:
                user = None

        if not user:
            # fallback
            user = fallback_users.get(str(user_id))

        if user:
            user_info = {
                "id": user_id,
                "email": data.get('email'),
                "name": user.get('name') or user.get('full_name') or '',
                "role": user.get('role', 'user')
            }
        else:
            user_info = {"id": user_id, "email": data.get('email'), "role": 'user'}

        return jsonify({
            "valid": True,
            "user": user_info
        }), 200
    except (jwt.ExpiredSignatureError, jwt.InvalidTokenError):
        return jsonify({"valid": False, "error": "Invalid or expired token"}), 401


# ===========================
# TRANSCRIBE AUDIO (PROTECTED) - WITH NOISE FILTERING
# ===========================
@app.route("/transcribe", methods=["POST"])
@token_required
def transcribe_audio(current_user_id):
    """Transcribe audio file with optional noise filtering (requires authentication)"""
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    audio_file = request.files["file"]
    apply_denoising = request.form.get("denoise", "false").lower() == "true"
    
    if audio_file.filename == "":
        return jsonify({"error": "Empty file"}), 400

    filepath = os.path.join(UPLOAD_FOLDER, audio_file.filename)
    audio_file.save(filepath)

    try:
        # Apply noise filtering if requested
        if apply_denoising:
            filepath = denoise_audio(filepath)

        # Lazy-load Whisper model if needed
        global model
        if model is None:
            try:
                import whisper
                model = whisper.load_model("base")
            except Exception as e:
                return jsonify({"error": f"Whisper model not available: {e}"}), 503

        # Transcribe with Whisper
        result = model.transcribe(filepath)
        text = result["text"]
        segments = extract_segments(result)

        # Save to MongoDB with user reference
        doc = {
            "user_id": str(current_user_id),
            "filename": audio_file.filename,
            "transcription": text,
            "segments": segments,
            "created_at": datetime.utcnow(),
            "summary": None,
            "key_items": None
        }
        transcription_result = transcriptions_collection.insert_one(doc)
        
        log_action("transcription_created", current_user_id, {"filename": audio_file.filename})
        track_metric("transcription_count", 1, str(current_user_id))

        return jsonify({
            "transcription_id": str(transcription_result.inserted_id),
            "transcription": text,
            "segments": segments
        }), 200
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ===========================
# TRANSLATE TEXT (PROTECTED)
# ===========================
@app.route("/translate", methods=["POST"])
@token_required
def translate_text(current_user_id):
    """Translate text to target language (requires authentication)"""
    data = request.get_json()
    text = data.get("text", "")
    target = data.get("target", "en")

    if not text:
        return jsonify({"error": "Text to translate is required"}), 400

    url = "https://translate.argosopentech.com/translate"

    payload = {
        "q": text,
        "source": "auto",
        "target": target,
        "format": "text"
    }

    try:
        response = requests.post(url, json=payload)
        result = response.json()
        
        log_action("translation", current_user_id, {"target_language": target})
        track_metric("translation_count", 1, str(current_user_id))
        
        return jsonify({"translatedText": result.get("translatedText", "")}), 200
    except Exception as e:
        return jsonify({"translatedText": "", "error": str(e)}), 500


# ===========================
# SUMMARIZATION (NEW FEATURE - AD-4)
# ===========================
@app.route("/transcriptions/<transcription_id>/summarize", methods=["POST"])
@token_required
def summarize_transcription(current_user_id, transcription_id):
    """Generate summary of transcription"""
    try:
        from bson import ObjectId
        
        transcription = transcriptions_collection.find_one({
            "_id": ObjectId(transcription_id),
            "user_id": str(current_user_id)
        })
        
        if not transcription:
            return jsonify({"error": "Transcription not found"}), 404
        
        # Lazy-load summarizer if needed
        global summarizer
        if summarizer is None:
            try:
                from transformers import pipeline
                summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
            except Exception as e:
                print(f"Summarizer could not be loaded: {e}")
                summarizer = None

        if not summarizer:
            return jsonify({"error": "Summarization service unavailable"}), 503
        
        text = transcription["transcription"]
        
        # Split text into sentences for better summarization
        sentences = text.split('. ')
        if len(sentences) < 3:
            return jsonify({"summary": text, "message": "Text too short to summarize"}), 200
        
        # Summarize
        summary = summarizer(text, max_length=150, min_length=30, do_sample=False)
        summary_text = summary[0]['summary_text']
        
        # Extract bullet points
        bullet_points = [s.strip() + "." for s in summary_text.split('.') if s.strip()]
        
        # Update transcription with summary
        transcriptions_collection.update_one(
            {"_id": ObjectId(transcription_id)},
            {"$set": {"summary": summary_text, "bullet_points": bullet_points}}
        )
        
        log_action("summarization", current_user_id, {"transcription_id": transcription_id})
        track_metric("summarization_count", 1, str(current_user_id))
        
        return jsonify({
            "summary": summary_text,
            "bullet_points": bullet_points
        }), 200
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ===========================
# KEY ITEMS ENDPOINTS (AD-6)
# ===========================
@app.route("/transcriptions/<transcription_id>/extract-items", methods=["POST"])
@token_required
def extract_items(current_user_id, transcription_id):
    """Extract key decisions and action items from a transcription"""
    try:
        from bson import ObjectId

        transcription = transcriptions_collection.find_one({
            "_id": ObjectId(transcription_id),
            "user_id": str(current_user_id)
        })

        if not transcription:
            return jsonify({"error": "Transcription not found"}), 404

        text = transcription.get('transcription', '')
        items = extract_key_items_from_text(text)

        # Update transcription
        transcriptions_collection.update_one({"_id": ObjectId(transcription_id)}, {"$set": {"key_items": items}})

        log_action("extract_key_items", current_user_id, {"transcription_id": transcription_id, "count": len(items)})
        track_metric("key_items_extracted", len(items), str(current_user_id))

        return jsonify({"key_items": items}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/transcriptions/<transcription_id>/key-items", methods=["POST"])
@token_required
def update_key_items(current_user_id, transcription_id):
    """Update key items array for a transcription (bulk replace). Accepts JSON { key_items: [...] }"""
    try:
        from bson import ObjectId
        data = request.get_json() or {}
        items = data.get('key_items')

        if items is None:
            return jsonify({"error": "key_items array required"}), 400

        # Basic validation: ensure list
        if not isinstance(items, list):
            return jsonify({"error": "key_items must be a list"}), 400

        result = transcriptions_collection.update_one(
            {"_id": ObjectId(transcription_id), "user_id": str(current_user_id)},
            {"$set": {"key_items": items}}
        )

        if result.matched_count == 0:
            return jsonify({"error": "Transcription not found or not owned by user"}), 404

        log_action("update_key_items", current_user_id, {"transcription_id": transcription_id, "count": len(items)})
        return jsonify({"message": "Key items updated"}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ===========================
# KEYWORD SEARCH (NEW FEATURE - Sprint 2 #4)
# ===========================
@app.route("/transcriptions/search", methods=["GET"])
@token_required
def search_transcriptions(current_user_id):
    """Search transcriptions by keyword"""
    query = request.args.get("q", "").strip()
    
    if not query:
        return jsonify({"error": "Search query required"}), 400
    
    # Create text index if not exists
    try:
        transcriptions_collection.create_index([("transcription", "text")])
    except:
        pass
    
    try:
        # Search with full-text search
        results = list(transcriptions_collection.find({
            "user_id": str(current_user_id),
            "$text": {"$search": query}
        }))
        
        # Format results
        for r in results:
            r['_id'] = str(r['_id'])
            r['created_at'] = r['created_at'].isoformat()
            
            # Find matching segments
            matches = []
            for segment in r.get('segments', []):
                if query.lower() in segment['text'].lower():
                    matches.append(segment)
            r['matching_segments'] = matches
        
        log_action("search", current_user_id, {"query": query, "results": len(results)})
        track_metric("search_count", 1, str(current_user_id))
        
        return jsonify({
            "query": query,
            "count": len(results),
            "results": results
        }), 200
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ===========================
# EXPORT TO PDF/DOCX (NEW FEATURE - Sprint 2 #7)
# ===========================
@app.route("/transcriptions/<transcription_id>/export", methods=["GET"])
@token_required
def export_transcription(current_user_id, transcription_id):
    """Export transcription as PDF or DOCX"""
    try:
        from bson import ObjectId
        
        format_type = request.args.get("format", "docx").lower()
        
        if format_type not in ["pdf", "docx"]:
            return jsonify({"error": "Format must be 'pdf' or 'docx'"}), 400
        
        transcription = transcriptions_collection.find_one({
            "_id": ObjectId(transcription_id),
            "user_id": str(current_user_id)
        })
        
        if not transcription:
            return jsonify({"error": "Transcription not found"}), 404
        
        # Create DOCX
        doc = Document()
        
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run("Meeting Minutes")
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Metadata
        doc.add_paragraph(f"Meeting: {transcription['filename']}")
        doc.add_paragraph(f"Date: {transcription['created_at'].strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        # Summary if available
        if transcription.get('summary'):
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(transcription['summary'])
            doc.add_paragraph()
        
        # Full Transcription
        doc.add_heading("Full Transcription", level=2)
        doc.add_paragraph(transcription['transcription'])
        
        # Segments with timestamps
        if transcription.get('segments'):
            doc.add_heading("Segments with Timestamps", level=2)
            for segment in transcription['segments']:
                start = segment.get('start', 0)
                end = segment.get('end', 0)
                text = segment.get('text', '')
                doc.add_paragraph(f"[{start:.2f}s - {end:.2f}s] {text}")
        
        # Bullet points if available
        if transcription.get('bullet_points'):
            doc.add_heading("Key Points", level=2)
            for point in transcription['bullet_points']:
                doc.add_paragraph(point, style='List Bullet')

        # Key items (decisions/action items)
        if transcription.get('key_items'):
            doc.add_heading("Decisions & Action Items", level=2)
            for ki in transcription['key_items']:
                assignee = ki.get('assignee') or ''
                status = ki.get('status') or 'open'
                text = ki.get('text') or ''
                line = f"[{status}] " + (f"{assignee}: " if assignee else "") + text
                doc.add_paragraph(line)
        
        # Save to bytes
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        log_action("export", current_user_id, {"transcription_id": transcription_id, "format": format_type})
        track_metric("export_count", 1, str(current_user_id))
        
        return send_file(
            doc_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"minutes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ===========================
# GET USER TRANSCRIPTIONS (PROTECTED)
# ===========================
@app.route("/transcriptions", methods=["GET"])
@token_required
def get_transcriptions(current_user_id):
    """Get all transcriptions for current user"""
    transcriptions = list(transcriptions_collection.find({"user_id": str(current_user_id)}))
    
    # Convert ObjectId to string for JSON serialization
    for t in transcriptions:
        t['_id'] = str(t['_id'])
        t['created_at'] = t['created_at'].isoformat()
    
    return jsonify({"transcriptions": transcriptions}), 200


# ===========================
# ADMIN ROUTES (NEW FEATURES - AD-11, AD-12, Sprint 2 #3)
# ===========================

@app.route("/admin/users", methods=["GET"])
@admin_required
def get_all_users(current_user_id):
    """Get all users (admin only)"""
    users = list(users_collection.find({}, {"password": 0}))
    
    for u in users:
        u['_id'] = str(u['_id'])
        u['created_at'] = u['created_at'].isoformat()
    
    log_action("admin_view_users", current_user_id)
    
    return jsonify({"users": users}), 200


@app.route("/admin/users/<user_id>", methods=["DELETE"])
@admin_required
def delete_user(current_user_id, user_id):
    """Delete a user (admin only)"""
    try:
        from bson import ObjectId
        
        if user_id == str(current_user_id):
            return jsonify({"error": "Cannot delete yourself"}), 400
        
        result = users_collection.delete_one({"_id": ObjectId(user_id)})
        
        if result.deleted_count == 0:
            return jsonify({"error": "User not found"}), 404
        
        log_action("admin_delete_user", current_user_id, {"deleted_user_id": user_id})
        
        return jsonify({"message": "User deleted successfully"}), 200
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/admin/logs", methods=["GET"])
@admin_required
def get_logs(current_user_id):
    """Get system logs (admin only)"""
    limit = int(request.args.get("limit", 100))
    action_filter = request.args.get("action", None)
    
    query = {}
    if action_filter:
        query["action"] = action_filter
    
    logs = list(logs_collection.find(query).sort("timestamp", -1).limit(limit))
    
    for log in logs:
        log['_id'] = str(log['_id'])
        log['timestamp'] = log['timestamp'].isoformat()
    
    log_action("admin_view_logs", current_user_id)
    
    return jsonify({"logs": logs}), 200


@app.route("/admin/analytics", methods=["GET"])
@admin_required
def get_analytics(current_user_id):
    """Get system analytics (admin only)"""
    try:
        # Get metrics
        total_users = users_collection.count_documents({})
        total_transcriptions = transcriptions_collection.count_documents({})
        total_logins = logs_collection.count_documents({"action": "user_login"})
        
        # Get transcriptions per user
        pipeline = [
            {"$group": {"_id": "$user_id", "count": {"$sum": 1}}},
            {"$sort": {"count": -1}},
            {"$limit": 10}
        ]
        top_users = list(transcriptions_collection.aggregate(pipeline))
        
        # Get metrics over time (last 7 days)
        seven_days_ago = datetime.utcnow() - timedelta(days=7)
        
        recent_metrics = list(analytics_collection.find({
            "timestamp": {"$gte": seven_days_ago}
        }).sort("timestamp", -1))
        
        for m in recent_metrics:
            m['_id'] = str(m['_id'])
            m['timestamp'] = m['timestamp'].isoformat()
        
        log_action("admin_view_analytics", current_user_id)
        
        return jsonify({
            "total_users": total_users,
            "total_transcriptions": total_transcriptions,
            "total_logins": total_logins,
            "top_users": top_users,
            "recent_metrics": recent_metrics
        }), 200
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ===========================
# HEALTH CHECK
# ===========================
@app.route("/", methods=["GET"])
def home():
    """Health check endpoint"""
    return jsonify({
        "message": "MinuteMinds backend is running!",
        "features": [
            "User Registration & Login",
            "Audio Transcription with Noise Filtering",
            "Automatic Summarization",
            "Keyword Search",
            "PDF/DOCX Export",
            "Admin Dashboard",
            "System Logs",
            "Analytics"
        ],
        "version": "2.0"
    }), 200


if __name__ == "__main__":
    app.run(port=5000, debug=True)
